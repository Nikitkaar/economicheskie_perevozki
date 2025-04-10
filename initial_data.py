import logging
import math
import re
from docx import Document


class WordEquationReplacer:
    def __init__(self, file_path, **kwargs):
        self.doc = Document(file_path)
        self.replacements = kwargs
        logging.debug(f"Инициализация WordEquationReplacer: {self.replacements}")

    import re

    def replace_text_in_paragraph(self, p):
        #Заменяет текст в переданном параграфе и логирует процесс.#
        for old, new in self.replacements.items():
            # Экранируем текст, чтобы избежать ошибок с регулярными выражениями
            sanitized_old = re.escape(old)

            # Используем re.sub для замены, просто указываем возможные разделители
            # Здесь мы обрабатываем возможные пробелы, запятые и точки
            pattern = re.compile(rf'\b{sanitized_old}\b')  # Слово должно быть целым, используем границы слова

            # Обработаем текст в параграфе
            if pattern.search(p.text):
                for run in p.runs:
                    if pattern.search(run.text):
                        old_text = run.text
                        run.text = pattern.sub(new, run.text)
                        logging.debug(f"Заменено в параграфе: '{old}' на '{new}' в тексте '{old_text}'")
            else:
                logging.debug(f"'{old}' не найдено в тексте: '{p.text}'")

    def replace_text_in_cell(self, cell):
        #Заменяет текст в переданной ячейке.#
        for paragraph in cell.paragraphs:
            self.replace_text_in_paragraph(paragraph)

    def process_document(self):
        #Обрабатывает документ, заменяя текст и обрабатывая таблицы.#
        logging.info("Обработка документа начата.")

        # Заменяем текст в параграфах
        for p in self.doc.paragraphs:
            logging.debug(f"Обрабатываем параграф: '{p.text}'")
            self.replace_text_in_paragraph(p)

        # Заменяем текст в таблицах
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    logging.debug(f"Обрабатываем ячейку: '{cell.text}'")
                    self.replace_text_in_cell(cell)

        logging.info("Обработка документа завершена.")

    def save_document(self, new_file_path):
        #Сохраняет документ в новый файл.
        self.doc.save(new_file_path)
        logging.info(f"Документ сохранен как: {new_file_path}")


class second_:
    # Статический словарь с данными о локомотивах
    locomotives_data = {
        "ВЛ60": {
            "Конструктивная скорость": 120,
            "Радиус колеса": 625,
            "Число осей": 3,
            "Длина жесткой базы": 4600,
            "Поперечные разбеги": {"крайних": 1.0, "средней у трехосной тележки": 15.5}
        },
        "ВЛ60к": {
            "Конструктивная скорость": 100,
            "Радиус колеса": 625,
            "Число осей": 3,
            "Длина жесткой базы": 4600,
            "Поперечные разбеги": {"крайних": 1.0, "средней у трехосной тележки": 15.5},
        },
        "ЧС4т": {
                "Конструктивная скорость": 160,
                "Радиус колеса": 625,
                "Число осей": 3,
                "Длина жесткой базы": 4600,
                "Поперечные разбеги": {"крайних": 1.3, "средней у трехосной тележки": 1.3}
                # Добавьте остальные локомотивы по аналогии
        }}

    def __init__(self, **kwargs):
        self.pepls_movie = kwargs.get("pepls_movie")
        self.townsman_up = kwargs.get("townsman_up")
        self.villager_up = kwargs.get("villager_up")
        self.villager_p = kwargs.get("villager_p")
        self.townsman_station = kwargs.get("townsman_station")

        self.wood = kwargs.get("wood")
        self.wood_products = kwargs.get("wood_products")
        self.Lumber = kwargs.get("Lumber")
        self.roundwood = kwargs.get("roundwood")
        self.firewood = kwargs.get("firewood")

        self.t = 10



        # Площади станций
        self.FctA = 894.14
        self.FctB = 686.41
        self.Fct = 894.1351 + 686.01

        # Население сельское на отчетный период
        self.AoVillA = self.villager_p * self.FctA
        self.AoVillB = self.villager_p * self.FctB

        # Сельское население на расчетный период
        self.ApVillA = self.AoVillA*((1+(self.villager_up/100))**self.t)
        self.ApVillB = self.AoVillB * ((1 + (self.villager_up/100)) ** self.t)

        # Городское население расчет
        # self.ApTownman = self.townsman_station*((1+(self.townsman_up/100))**self.t)
        self.ApTownman = self.townsman_station * ((1 + self.townsman_up/100) ** 10)

        # Всего населения на станции
        self.pepleA = self.ApVillA + (self.ApTownman*1000)
        self.pepleB = self.ApVillB

        # Лесосырье
        self.true_wood = 0.7 * self.wood
        self.true_wood_products = self.true_wood * (self.wood_products/100)
        self.true_Lumber = self.true_wood * (self.Lumber/100)
        self.true_roundwood = self.true_wood * (self.roundwood/100)
        self.true_firewood = self.true_wood * (self.firewood/100)
        # ВСЕГО
        self.all = (self.true_wood_products + self.true_Lumber + self.true_roundwood + self.true_firewood) * 1000

class third_():
    def __init__(self, **kwargs):
            self.FctB = kwargs.get("FctB")
            self.FctA = kwargs.get("FctA")
            self.plot_posevov = kwargs.get("plot_posevov")
            self.ves_posevov = kwargs.get("ves_posevov")
            self.up_posevov = kwargs.get("up_posevov")
            self.eat_posevov = kwargs.get("eat_posevov")
            self.need_city_wheat = kwargs.get("need_city_wheat")
            self.need_vilolage_wheat = kwargs.get("need_vilolage_wheat")
            self.ves_furaj = kwargs.get("ves_furaj")
            self.up_furaj = kwargs.get("up_furaj")
            self.eat_furaj = kwargs.get("eat_furaj")
            self.ves_potato = kwargs.get("ves_potato")
            self.up_potato = kwargs.get("up_potato")
            self.eat_potato = kwargs.get("eat_potato")
            self.need_city_potato = kwargs.get("need_city_potato")
            self.need_village_potato = kwargs.get("need_village_potato")
            self.ves_culture = kwargs.get("ves_culture")
            self.up_culture = kwargs.get("up_culture")

            # Площади станций
            #self.FctA = 894.1351
            #self.FctB = 686.4126
            #self.Fct = 894.1351 + 686.4126

            # Общая посевная площадь
            self.Fct_pos_A = self.FctA * self.plot_posevov / 1000
            self.Fct_pos_B = self.FctB * self.plot_posevov / 1000










